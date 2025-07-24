import streamlit as st
import pandas as pd
import os
from io import BytesIO
from PIL import Image
import fitz  # PyMuPDF
import pymupdf4llm
from typing import List, Dict, Any
import tempfile
import base64
from groq import Groq
from app.utils.config import Config
import time
import gc
import json

class PureLLMFileProcessor:
    """Pure LLM-based file processor - No regex, no hardcoding, LLMs do everything"""
    
    def __init__(self, groq_api_key: str):
        self.groq_client = Groq(api_key=groq_api_key)
        self.vision_model = Config.VISION_MODEL
        self.text_model = Config.TEXT_MODEL
        Config.ensure_directories()
        Config.cleanup_temp_directory()
        
        self.supported_formats = {
            'pdf': self.process_pdf_with_heavy_vision,
            'csv': self.process_csv_with_llm,
            'xlsx': self.process_excel_with_llm,
            'xls': self.process_excel_with_llm,
            'txt': self.process_text_with_llm,
            'docx': self.process_docx_with_llm,
            'png': self.process_image_with_vision,
            'jpg': self.process_image_with_vision,
            'jpeg': self.process_image_with_vision,
            'las': self.process_las_with_llm,
            'tiff': self.process_tiff_with_vision,
            'tif': self.process_tiff_with_vision
        }
    
    def encode_image_to_base64(self, image_data: bytes) -> str:
        """Convert image bytes to base64 string"""
        return base64.b64encode(image_data).decode('utf-8')
    
    def advanced_vision_analysis(self, image_data: bytes, context: str = "", analysis_type: str = "comprehensive") -> Dict[str, Any]:
        """Advanced LLM vision analysis with specialized prompting"""
        try:
            base64_image = self.encode_image_to_base64(image_data)
            
            # Advanced geological vision prompt
            vision_prompt = f"""
            You are an expert geological analyst with decades of experience in well log interpretation, formation analysis, and petroleum geology.
            
            Analyze this geological document image with EXTREME DETAIL and PRECISION. Extract EVERY piece of information visible.
            
            CRITICAL ANALYSIS REQUIREMENTS:
            
            1. **WELL IDENTIFICATION DATA**:
               - Extract ALL well names, API numbers, permit numbers
               - Identify operator/company names with complete details
               - Find location information (county, state, country, coordinates)
               - Extract lease names, section/township/range data
               
            2. **TECHNICAL WELL DATA**:
               - Total depth (TD), measured depth (MD), true vertical depth (TVD)
               - Spud date, completion date, first production date
               - Well type (vertical, horizontal, deviated)
               - Casing programs, completion details
               
            3. **FORMATION & GEOLOGICAL DATA**:
               - Formation names and tops with exact depths
               - Lithology descriptions and interpretations
               - Stratigraphic correlations and sequences
               - Geological age assignments
               
            4. **LOG & PETROPHYSICAL DATA**:
               - Log types run (GR, SP, RHOB, NPHI, RT, etc.)
               - Log dates and service companies
               - Porosity, permeability, water saturation values
               - Net-to-gross ratios, pay zone identification
               
            5. **DRILLING & COMPLETION DATA**:
               - Drilling fluid types and properties
               - Bit sizes and casing schedules
               - Cement volumes and placement
               - Perforation intervals and completion methods
               
            6. **PRODUCTION & TESTING DATA**:
               - Flow rates (oil, gas, water)
               - Pressure data (formation, wellhead, tubing)
               - Test results and interpretations
               - Production history if available
               
            7. **VISUAL ELEMENTS**:
               - Charts, graphs, and curve interpretations
               - Tables with numerical data
               - Diagrams and schematics
               - Annotations and handwritten notes
            
            Context: {context}
            Analysis Type: {analysis_type}
            
            RESPONSE FORMAT:
            Provide a comprehensive JSON-like structured analysis with:
            - Clear categorization of all extracted data
            - Exact values with units where applicable
            - Confidence levels for each piece of information
            - Relationships between different data points
            - Any uncertainties or ambiguities noted
            
            Be exhaustive and precise. Extract EVERYTHING visible, no matter how small or seemingly insignificant.
            """
            
            response = self.groq_client.chat.completions.create(
                model=self.vision_model,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": vision_prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                    ]
                }],
                temperature=0.05,  # Very low temperature for precision
                max_tokens=4000    # Maximum tokens for comprehensive analysis
            )
            
            return {
                'analysis': response.choices[0].message.content,
                'success': True,
                'analysis_type': analysis_type
            }
            
        except Exception as e:
            return {
                'analysis': f"Advanced vision analysis failed: {str(e)}",
                'success': False,
                'analysis_type': analysis_type
            }
    
    def llm_text_analysis(self, text_content: str, filename: str, content_type: str = "geological_document") -> Dict[str, Any]:
        """Advanced LLM text analysis for comprehensive data extraction"""
        try:
            analysis_prompt = f"""
            You are a world-class geological data analyst and petroleum engineer with expertise in well log interpretation, formation analysis, and geological data extraction.
            
            Analyze the following geological document text with MAXIMUM PRECISION and extract ALL relevant information.
            
            Document: {filename}
            Content Type: {content_type}
            
            COMPREHENSIVE EXTRACTION REQUIREMENTS:
            
            1. **WELL IDENTIFICATION**:
               - Well names (all variations and aliases)
               - API numbers and permit numbers
               - Operator and company information
               - Location data (coordinates, county, state, legal descriptions)
               
            2. **TECHNICAL SPECIFICATIONS**:
               - Depths (total depth, measured depth, TVD)
               - Dates (spud, completion, first production, log dates)
               - Well trajectory and type
               - Casing and completion details
               
            3. **GEOLOGICAL INFORMATION**:
               - Formation names and stratigraphic units
               - Formation tops with precise depths
               - Lithological descriptions
               - Geological age and depositional environment
               
            4. **PETROPHYSICAL DATA**:
               - Log types and service companies
               - Porosity, permeability, saturation values
               - Net pay calculations
               - Reservoir quality assessments
               
            5. **OPERATIONAL DATA**:
               - Drilling parameters and mud properties
               - Testing results and flow rates
               - Production data and completion methods
               - Equipment and service company information
               
            6. **QUANTITATIVE DATA**:
               - All numerical values with proper units
               - Ranges and intervals
               - Calculated parameters
               - Statistical summaries
            
            TEXT TO ANALYZE:
            {text_content[:15000]}  # Limit to avoid token limits
            
            RESPONSE REQUIREMENTS:
            - Provide structured, comprehensive extraction
            - Include confidence levels for extracted data
            - Identify relationships between data points
            - Note any inconsistencies or ambiguities
            - Format as clear, organized information blocks
            
            Extract EVERYTHING relevant with maximum detail and precision. Leave nothing behind.
            """
            
            response = self.groq_client.chat.completions.create(
                model=self.text_model,
                messages=[
                    {"role": "user", "content": analysis_prompt}
                ],
                temperature=0.1,
                max_tokens=4000
            )
            
            return {
                'analysis': response.choices[0].message.content,
                'success': True,
                'content_type': content_type
            }
            
        except Exception as e:
            return {
                'analysis': f"LLM text analysis failed: {str(e)}",
                'success': False,
                'content_type': content_type
            }
    
    def process_pdf_with_heavy_vision(self, file_bytes: bytes, filename: str, force_vision: bool = True) -> Dict[str, Any]:
        """Heavy vision processing for PDFs - analyze EVERY page with vision models"""
        pdf_document = None
        temp_file_path = None
        
        try:
            # Open PDF from memory
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
            
            # Extract text using pymupdf4llm ONLY
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf', dir=Config.TEMP_DIR) as temp_file:
                temp_file.write(file_bytes)
                temp_file_path = temp_file.name
            
            # Get structured text
            text_data = pymupdf4llm.to_markdown(temp_file_path)
            
            # Clean up temp file immediately
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except:
                    st.warning(f"Could not delete temp file: {temp_file_path}")
            
            # LLM text analysis
            text_analysis = self.llm_text_analysis(text_data, filename, "pdf_document")
            
            # HEAVY VISION ANALYSIS - Process EVERY page
            vision_analyses = []
            page_summaries = []
            
            total_pages = len(pdf_document)
            st.info(f"🔍 Performing heavy vision analysis on {total_pages} pages...")
            
            for page_num in range(total_pages):
                try:
                    page = pdf_document.load_page(page_num)
                    
                    # Convert page to high-quality image
                    pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))  # High resolution
                    page_image_data = pix.tobytes("png")
                    pix = None
                    
                    # Comprehensive vision analysis for each page
                    page_analysis = self.advanced_vision_analysis(
                        page_image_data, 
                        f"Page {page_num + 1} of {total_pages} from geological document {filename}",
                        "page_comprehensive"
                    )
                    
                    vision_analyses.append({
                        'page': page_num + 1,
                        'analysis': page_analysis['analysis'],
                        'success': page_analysis['success']
                    })
                    
                    # Create page summary
                    if page_analysis['success']:
                        page_summaries.append(f"Page {page_num + 1}: {page_analysis['analysis'][:500]}...")
                    
                    # Rate limiting
                    time.sleep(0.3)
                    
                    # Progress update
                    if (page_num + 1) % 5 == 0:
                        st.info(f"Processed {page_num + 1}/{total_pages} pages with vision analysis")
                    
                except Exception as e:
                    st.warning(f"Error processing page {page_num + 1}: {e}")
                    vision_analyses.append({
                        'page': page_num + 1,
                        'analysis': f"Error processing page: {str(e)}",
                        'success': False
                    })
            
            # Combine all analyses using LLM synthesis
            combined_analysis = self.synthesize_multimodal_analysis(
                text_analysis, vision_analyses, filename
            )
            
            return {
                'text': combined_analysis,
                'raw_text': text_data,
                'text_analysis': text_analysis,
                'vision_analyses': vision_analyses,
                'pages': total_pages,
                'processing_stats': {
                    'vision_calls_made': len(vision_analyses),
                    'processing_type': 'heavy_vision_llm',
                    'text_analysis': 'advanced_llm',
                    'vision_analysis': 'comprehensive_per_page'
                },
                'metadata': {
                    'filename': filename,
                    'type': 'pdf_heavy_vision_llm',
                    'has_vision_analysis': True,
                    'has_text_analysis': True,
                    'processing_approach': 'pure_llm_multimodal'
                }
            }
            
        except Exception as e:
            return {
                'text': f"Error processing PDF {filename}: {str(e)}", 
                'metadata': {'filename': filename, 'type': 'pdf_heavy_vision', 'error': True}
            }
        finally:
            if pdf_document:
                pdf_document.close()
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
            gc.collect()
    
    def synthesize_multimodal_analysis(self, text_analysis: Dict, vision_analyses: List[Dict], filename: str) -> str:
        """LLM-based synthesis of text and vision analyses"""
        try:
            # Prepare synthesis data
            text_content = text_analysis.get('analysis', 'No text analysis available')
            
            vision_content = "\n\n".join([
                f"Page {va['page']}: {va['analysis']}" 
                for va in vision_analyses if va['success']
            ])
            
            synthesis_prompt = f"""
            You are an expert geological analyst tasked with creating a comprehensive, unified analysis by synthesizing multiple data sources.
            
            You have been provided with:
            1. Advanced text analysis of document: {filename}
            2. Comprehensive vision analysis of all pages
            
            SYNTHESIS REQUIREMENTS:
            Create a unified, comprehensive geological document analysis that:
            
            1. **CONSOLIDATES ALL INFORMATION** from both text and vision analyses
            2. **RESOLVES CONFLICTS** between different sources intelligently
            3. **FILLS GAPS** where one analysis complements the other
            4. **STRUCTURES INFORMATION** in a logical, hierarchical format
            5. **MAINTAINS PRECISION** of all technical data and measurements
            
            TEXT ANALYSIS RESULTS:
            {text_content}
            
            VISION ANALYSIS RESULTS:
            {vision_content}
            
            FINAL SYNTHESIS FORMAT:
            Create a comprehensive document that includes:
            
            # GEOLOGICAL DOCUMENT ANALYSIS: {filename}
            
            ## EXECUTIVE SUMMARY
            [Brief overview of the document and key findings]
            
            ## WELL IDENTIFICATION
            [All well identifiers, names, API numbers, locations]
            
            ## TECHNICAL SPECIFICATIONS
            [Depths, dates, well type, completion details]
            
            ## GEOLOGICAL DATA
            [Formations, lithology, stratigraphy, age]
            
            ## PETROPHYSICAL ANALYSIS
            [Logs, porosity, permeability, saturation, net pay]
            
            ## OPERATIONAL INFORMATION
            [Drilling, completion, testing, production]
            
            ## QUANTITATIVE DATA SUMMARY
            [All numerical values, ranges, calculations]
            
            ## VISUAL ELEMENTS INTERPRETATION
            [Charts, graphs, diagrams, tables from vision analysis]
            
            ## INTEGRATED CONCLUSIONS
            [Key insights from combined text and vision analysis]
            
            Be exhaustive, precise, and comprehensive. This should be the definitive analysis of this geological document.
            """
            
            response = self.groq_client.chat.completions.create(
                model=self.text_model,
                messages=[
                    {"role": "user", "content": synthesis_prompt}
                ],
                temperature=0.1,
                max_tokens=4000
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            # Fallback synthesis
            return f"""
            # GEOLOGICAL DOCUMENT ANALYSIS: {filename}
            
            ## TEXT ANALYSIS:
            {text_analysis.get('analysis', 'Text analysis failed')}
            
            ## VISION ANALYSIS:
            {vision_content if 'vision_content' in locals() else 'Vision analysis compilation failed'}
            
            ## SYNTHESIS ERROR:
            LLM synthesis failed: {str(e)}
            """
    
    def process_csv_with_llm(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """LLM-based CSV analysis"""
        try:
            df = pd.read_csv(BytesIO(file_bytes))
            
            # Convert to text representation
            csv_text = f"CSV File: {filename}\n\n"
            csv_text += f"Shape: {df.shape}\n"
            csv_text += f"Columns: {', '.join(df.columns)}\n\n"
            csv_text += "Data Sample:\n"
            csv_text += df.to_string()
            
            # LLM analysis
            analysis = self.llm_text_analysis(csv_text, filename, "csv_data")
            
            return {
                'text': analysis['analysis'],
                'raw_data': csv_text,
                'metadata': {
                    'filename': filename,
                    'type': 'csv_llm_analyzed',
                    'rows': len(df),
                    'columns': len(df.columns),
                    'analysis_method': 'pure_llm'
                }
            }
        except Exception as e:
            return {'text': f"Error processing CSV {filename}: {str(e)}", 'metadata': {'filename': filename, 'type': 'csv', 'error': True}}
    
    def process_excel_with_llm(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """LLM-based Excel analysis"""
        try:
            excel_file = pd.ExcelFile(BytesIO(file_bytes))
            
            excel_text = f"Excel File: {filename}\n\nSheets: {', '.join(excel_file.sheet_names)}\n\n"
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                excel_text += f"Sheet '{sheet_name}':\n{df.to_string()}\n\n"
            
            # LLM analysis
            analysis = self.llm_text_analysis(excel_text, filename, "excel_data")
            
            return {
                'text': analysis['analysis'],
                'raw_data': excel_text,
                'metadata': {
                    'filename': filename,
                    'type': 'excel_llm_analyzed',
                    'sheets': len(excel_file.sheet_names),
                    'analysis_method': 'pure_llm'
                }
            }
        except Exception as e:
            return {'text': f"Error processing Excel {filename}: {str(e)}", 'metadata': {'filename': filename, 'type': 'excel', 'error': True}}
    
    def process_text_with_llm(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """LLM-based text analysis"""
        try:
            text = file_bytes.decode('utf-8')
            analysis = self.llm_text_analysis(text, filename, "text_document")
            
            return {
                'text': analysis['analysis'],
                'raw_text': text,
                'metadata': {
                    'filename': filename,
                    'type': 'text_llm_analyzed',
                    'length': len(text),
                    'analysis_method': 'pure_llm'
                }
            }
        except Exception as e:
            return {'text': f"Error processing text {filename}: {str(e)}", 'metadata': {'filename': filename, 'type': 'text', 'error': True}}
    
    def process_docx_with_llm(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """LLM-based DOCX analysis"""
        try:
            from docx import Document
            doc = Document(BytesIO(file_bytes))
            
            docx_text = f"Document: {filename}\n\n"
            for paragraph in doc.paragraphs:
                docx_text += paragraph.text + "\n"
            
            analysis = self.llm_text_analysis(docx_text, filename, "docx_document")
            
            return {
                'text': analysis['analysis'],
                'raw_text': docx_text,
                'metadata': {
                    'filename': filename,
                    'type': 'docx_llm_analyzed',
                    'paragraphs': len(doc.paragraphs),
                    'analysis_method': 'pure_llm'
                }
            }
        except Exception as e:
            return {'text': f"Error processing DOCX {filename}: {str(e)}", 'metadata': {'filename': filename, 'type': 'docx', 'error': True}}
    
    def process_las_with_llm(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """LLM-based LAS file analysis"""
        try:
            las_content = file_bytes.decode('utf-8', errors='ignore')
            analysis = self.llm_text_analysis(las_content, filename, "las_well_log")
            
            return {
                'text': analysis['analysis'],
                'raw_las': las_content,
                'metadata': {
                    'filename': filename,
                    'type': 'las_llm_analyzed',
                    'analysis_method': 'pure_llm',
                    'has_geological_data': True
                }
            }
        except Exception as e:
            return {'text': f"Error processing LAS {filename}: {str(e)}", 'metadata': {'filename': filename, 'type': 'las', 'error': True}}
    
    def process_image_with_vision(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Heavy vision analysis for images"""
        try:
            image = Image.open(BytesIO(file_bytes))
            
            # Comprehensive vision analysis
            vision_analysis = self.advanced_vision_analysis(
                file_bytes, 
                f"Geological image file: {filename}",
                "comprehensive_image"
            )
            
            return {
                'text': vision_analysis['analysis'],
                'metadata': {
                    'filename': filename,
                    'type': 'image_heavy_vision',
                    'size': image.size,
                    'mode': image.mode,
                    'has_vision_analysis': vision_analysis['success'],
                    'analysis_method': 'advanced_vision_llm'
                }
            }
        except Exception as e:
            return {'text': f"Error processing image {filename}: {str(e)}", 'metadata': {'filename': filename, 'type': 'image', 'error': True}}
    
    def process_tiff_with_vision(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Heavy vision analysis for TIFF images"""
        try:
            image = Image.open(BytesIO(file_bytes))
            
            # Convert to RGB if needed
            if image.mode not in ['RGB', 'RGBA']:
                image = image.convert('RGB')
            
            # Convert to PNG for vision processing
            png_buffer = BytesIO()
            image.save(png_buffer, format='PNG')
            png_bytes = png_buffer.getvalue()
            
            # Comprehensive vision analysis
            vision_analysis = self.advanced_vision_analysis(
                png_bytes, 
                f"TIFF geological image: {filename}",
                "comprehensive_tiff"
            )
            
            return {
                'text': vision_analysis['analysis'],
                'metadata': {
                    'filename': filename,
                    'type': 'tiff_heavy_vision',
                    'size': image.size,
                    'mode': image.mode,
                    'has_vision_analysis': vision_analysis['success'],
                    'analysis_method': 'advanced_vision_llm'
                }
            }
        except Exception as e:
            return {'text': f"Error processing TIFF {filename}: {str(e)}", 'metadata': {'filename': filename, 'type': 'tiff', 'error': True}}
    
    def process_file(self, uploaded_file, force_vision: bool = True) -> Dict[str, Any]:
        """Process uploaded file with pure LLM approach"""
        filename = uploaded_file.name
        file_extension = filename.split('.')[-1].lower()
        
        # Check file size
        if hasattr(uploaded_file, 'size') and uploaded_file.size > Config.MAX_FILE_SIZE_MB * 1024 * 1024:
            return {
                'text': f"File {filename} exceeds {Config.MAX_FILE_SIZE_MB}MB limit.",
                'metadata': {'filename': filename, 'type': 'error', 'error': True}
            }
        
        try:
            file_bytes = uploaded_file.read()
        except Exception as e:
            return {
                'text': f"Error reading file {filename}: {str(e)}",
                'metadata': {'filename': filename, 'type': 'error', 'error': True}
            }
        
        if file_extension in self.supported_formats:
            try:
                return self.supported_formats[file_extension](file_bytes, filename)
            except Exception as e:
                return {
                    'text': f"Error processing {filename}: {str(e)}",
                    'metadata': {'filename': filename, 'type': file_extension, 'error': True}
                }
        else:
            return {
                'text': f"Unsupported file format: {file_extension}",
                'metadata': {'filename': filename, 'type': 'unsupported', 'error': True}
            }
